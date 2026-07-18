"""重新產生 tests/fixtures/docs/ 的樣本檔(T2.2)。

二進位樣本(pdf / docx)無法以 diff 審查,故保留產生腳本作為其「原始碼」:
需要調整樣本內容時改這裡並重跑,NEVER 手工替換不明來源的檔案。

    python tests/fixtures/docs/make_fixtures.py
"""
from pathlib import Path

import pymupdf
from docx import Document

HERE = Path(__file__).parent

TXT = """第一段落,包含全形空白　與一般空白。

第二段落。



第三段落結束。
"""

MD = """# 標題一

前言段落。

## 標題二

- 項目一
- 項目二

```python
print("hello")
```

結尾段落。
"""

HTML = """<!doctype html>
<html><head><title>t</title>
<style>body { color: red; }</style>
<script>console.log("noise");</script>
</head>
<body>
<h1>主標題</h1>
<p>第一段內文。</p>
<h2>次標題</h2>
<ul><li>清單項目一</li><li>清單項目二</li></ul>
<table>
  <tr><th>欄位</th><th>數值</th></tr>
  <tr><td>甲</td><td>1</td></tr>
  <tr><td><p>乙</p></td><td>2</td></tr>
</table>
<p>   </p>
</body></html>
"""


def write_text_fixtures() -> None:
    (HERE / "sample.txt").write_bytes(TXT.encode("utf-8"))
    (HERE / "sample_big5.txt").write_bytes("繁體中文 Big5 內容。\n\n第二段。\n".encode("cp950"))
    (HERE / "sample.md").write_bytes(MD.encode("utf-8"))
    (HERE / "sample.html").write_bytes(HTML.encode("utf-8"))


def write_docx_fixture() -> None:
    document = Document()
    document.add_heading("報告標題", level=1)
    document.add_paragraph("這是一段內文。")
    document.add_heading("章節標題", level=2)
    document.add_paragraph("清單項目一", style="List Bullet")
    document.add_paragraph("清單項目二", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "欄位"
    table.cell(0, 1).text = "數值"
    table.cell(1, 0).text = "甲"
    table.cell(1, 1).text = "100"
    document.add_paragraph("")  # 空段落:應被共同規則剔除
    document.save(str(HERE / "sample.docx"))


def _build_pdf() -> pymupdf.Document:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 90), "文件主標題", fontsize=24, fontname="china-t")
    page.insert_text((72, 140), "這是第一段內文,字級與標題明顯不同。", fontsize=11,
                     fontname="china-t")
    page.insert_text((72, 170), "這是第二段內文,用於驗證 block 切分。", fontsize=11,
                     fontname="china-t")
    return document


def write_pdf_fixtures() -> None:
    with _build_pdf() as document:
        document.save(HERE / "sample.pdf")
    with _build_pdf() as document:
        document.save(
            HERE / "encrypted.pdf",
            encryption=pymupdf.PDF_ENCRYPT_AES_256,  # type: ignore[attr-defined]  # pymupdf stub 未收錄
            owner_pw="owner-secret",
            user_pw="user-secret",
        )


if __name__ == "__main__":
    write_text_fixtures()
    write_docx_fixture()
    write_pdf_fixtures()
    for path in sorted(HERE.iterdir()):
        if path.name != Path(__file__).name:
            print(f"{path.name}: {path.stat().st_size} bytes")
